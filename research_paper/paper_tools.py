#!/usr/bin/env python3
"""
Research Paper Tools - A unified script for working with the research paper.

This script provides various functions for working with the research paper:
1. Combining all markdown sections into a single file
2. Converting to different formats (PDF, DOCX)
3. Creating reference templates

Usage:
  python paper_tools.py combine   # Combine all sections into complete_paper.md
  python paper_tools.py docx      # Create a Word document (log_analyzer_paper.docx)
  python paper_tools.py pdf       # Create a PDF document (requires pandoc)
  python paper_tools.py all       # Perform all conversions
  python paper_tools.py clean     # Remove temporary files
"""

import os
import re
import sys
import shutil
import subprocess
from datetime import datetime

# Check if python-docx is installed
DOCX_AVAILABLE = False
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    pass

# Check if pandoc is installed
PANDOC_AVAILABLE = False
PANDOC_PATH = "C:\\Users\\ameya\\AppData\\Local\\Pandoc\\pandoc.exe"
try:
    subprocess.run([PANDOC_PATH, "--version"], capture_output=True, check=True)
    PANDOC_AVAILABLE = True
except (subprocess.SubprocessError, FileNotFoundError):
    try:
        subprocess.run(["pandoc", "--version"], capture_output=True, check=True)
        PANDOC_PATH = "pandoc"
        PANDOC_AVAILABLE = True
    except (subprocess.SubprocessError, FileNotFoundError):
        pass

def read_file(file_path):
    """Read a file and return its content as a string."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        print(f"Error reading file {file_path}: {e}")
        return f"[Error: Could not read {file_path}]"

def process_includes(content, base_dir):
    """Process [Include: filename.md] tags and replace with file content."""
    include_pattern = r'\[Include: ([^\]]+)\]'

    def replace_include(match):
        filename = match.group(1).strip()
        file_path = os.path.join(base_dir, filename)
        return read_file(file_path)

    return re.sub(include_pattern, replace_include, content)

def combine_sections(output_file="complete_paper_updated.md"):
    """Update the complete paper with latest section content."""
    print(f"Updating complete paper from individual sections...")

    # List of section files in order
    section_files = [
        "abstract_introduction.md",
        "related_work.md",
        "system_architecture.md",
        "log_format_support.md",
        "remote_log_acquisition.md",
        "data_processing.md",
        "visualization_techniques.md",
        "case_studies.md",
        "performance_evaluation.md",
        "methodology_validation.md",
        "future_work.md",
        "conclusion.md",
        "references.md",
        "appendices.md"
    ]

    try:
        combined_content = []
        combined_content.append("# Log Analyzer: A Versatile Framework for Multi-Format Log Analysis in Cybersecurity Applications\n")
        combined_content.append("*Complete research paper with empirical validation*\n\n")

        for section_file in section_files:
            if os.path.exists(section_file):
                print(f"Adding section: {section_file}")
                content = read_file(section_file)
                combined_content.append("---\n\n")
                combined_content.append(content)
                combined_content.append("\n\n")
            else:
                print(f"Warning: Section file not found: {section_file}")

        with open(output_file, 'w', encoding='utf-8') as file:
            file.write(''.join(combined_content))
        print(f"Successfully updated complete paper: {output_file}")
        return True
    except Exception as e:
        print(f"Error writing output file {output_file}: {e}")
        return False

def create_docx_with_library(input_file="complete_paper.md", output_file="log_analyzer_paper.docx"):
    """Create a Word document using python-docx library."""
    if not DOCX_AVAILABLE:
        print("Error: python-docx library is not installed.")
        print("Please install it using: pip install python-docx")
        return False

    print(f"Creating Word document from {input_file} to {output_file}...")

    base_dir = os.path.dirname(os.path.abspath(input_file))
    content = read_file(input_file)
    processed_content = process_includes(content, base_dir)

    # Create a new Document
    doc = Document()

    # Set document properties
    doc.core_properties.title = "Log Analyzer: A Versatile Framework for Multi-Format Log Analysis in Cybersecurity Applications"
    doc.core_properties.author = "Research Team"
    doc.core_properties.created = datetime.now()

    # Add title
    title = doc.add_heading("Log Analyzer: A Versatile Framework for Multi-Format Log Analysis in Cybersecurity Applications", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Process markdown content
    sections = processed_content.split("---")
    for section in sections:
        if not section.strip():
            continue

        # Add section content
        doc.add_paragraph(section.strip())
        doc.add_page_break()

    # Save the document
    try:
        doc.save(output_file)
        print(f"Successfully created Word document: {output_file}")
        return True
    except Exception as e:
        print(f"Error saving document: {e}")
        return False

def create_docx_with_pandoc(input_file="complete_paper.md", output_file="log_analyzer_paper.docx"):
    """Create a Word document using pandoc."""
    if not PANDOC_AVAILABLE:
        print("Error: Pandoc is not installed or not in PATH.")
        print("Please install pandoc from https://pandoc.org/installing.html")
        return False

    print(f"Creating Word document using pandoc from {input_file} to {output_file}...")

    # Ensure the input file exists
    if not os.path.exists(input_file):
        # Try to combine sections first
        if not combine_sections(output_file=input_file):
            return False

    try:
        # Run pandoc to convert to docx
        cmd = [
            PANDOC_PATH,
            input_file,
            "-o", output_file,
            "--toc",
            "--toc-depth=3"
        ]

        print(f"Running command: {' '.join(cmd)}")
        subprocess.run(cmd, check=True)
        print(f"Successfully created Word document: {output_file}")
        return True
    except Exception as e:
        print(f"Error creating Word document: {e}")
        return False

def create_pdf_with_pandoc(input_file="complete_paper.md", output_file="log_analyzer_paper.pdf"):
    """Create a PDF document using pandoc."""
    if not PANDOC_AVAILABLE:
        print("Error: Pandoc is not installed or not in PATH.")
        print("Please install pandoc from https://pandoc.org/installing.html")
        return False

    print(f"Creating PDF document using pandoc from {input_file} to {output_file}...")

    # Ensure the input file exists
    if not os.path.exists(input_file):
        # Try to combine sections first
        if not combine_sections(output_file=input_file):
            return False

    try:
        # Run pandoc to convert to PDF
        # Try without specifying a PDF engine first
        cmd = [
            PANDOC_PATH,
            input_file,
            "-o", output_file,
            "--toc",
            "--toc-depth=3"
        ]

        print(f"Running command: {' '.join(cmd)}")
        subprocess.run(cmd, check=True)
        print(f"Successfully created PDF document: {output_file}")
        return True
    except Exception as e:
        print(f"Error creating PDF document: {e}")
        return False

def clean_temporary_files():
    """Remove temporary files."""
    print("Cleaning temporary files...")

    files_to_remove = [
        "complete_paper.md",
        "__pycache__"
    ]

    for file_path in files_to_remove:
        if os.path.exists(file_path):
            try:
                if os.path.isdir(file_path):
                    shutil.rmtree(file_path)
                else:
                    os.remove(file_path)
                print(f"Removed: {file_path}")
            except Exception as e:
                print(f"Error removing {file_path}: {e}")

    print("Cleanup complete.")
    return True

def print_help():
    """Print help information."""
    print(__doc__)
    print("\nAvailable formats:")
    print("  - DOCX: " + ("Available" if DOCX_AVAILABLE else "Not available (install python-docx)"))
    print("  - PDF: " + ("Available" if PANDOC_AVAILABLE else "Not available (install pandoc)"))

def main():
    """Main function to parse arguments and perform actions."""
    if len(sys.argv) < 2:
        print_help()
        return

    command = sys.argv[1].lower()

    if command == "combine":
        combine_sections()
    elif command == "docx":
        if DOCX_AVAILABLE:
            create_docx_with_library()
        elif PANDOC_AVAILABLE:
            create_docx_with_pandoc()
        else:
            print("Error: Neither python-docx nor pandoc is available.")
            print("Please install one of them to create Word documents.")
    elif command == "pdf":
        create_pdf_with_pandoc()
    elif command == "all":
        combine_sections()
        if DOCX_AVAILABLE:
            create_docx_with_library()
        elif PANDOC_AVAILABLE:
            create_docx_with_pandoc()
        create_pdf_with_pandoc()
    elif command == "clean":
        clean_temporary_files()
    elif command in ["help", "-h", "--help"]:
        print_help()
    else:
        print(f"Unknown command: {command}")
        print_help()

if __name__ == "__main__":
    main()
